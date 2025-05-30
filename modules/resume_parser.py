"""
Resume Parser Module - Extracts structured data from resume DOCX files.
"""
import json
import docx
import re

def extract_text_from_docx(docx_path):
    """
    Extract all text content from a DOCX file.
    
    Args:
        docx_path (str): Path to the DOCX file
        
    Returns:
        str: Full text content of the document
    """
    doc = docx.Document(docx_path)
    full_text = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text.append(paragraph.text)
    
    return '\n'.join(full_text)

def identify_sections(text):
    """
    Attempt to identify resume sections based on common headings.
    
    Args:
        text (str): Full resume text
        
    Returns:
        dict: Dictionary with identified sections
    """
    # Common section headers in resumes
    section_patterns = {
        'contact': r'(?i)(personal\s+information|contact|contact\s+information)',
        'summary': r'(?i)(summary|professional\s+summary|profile|objective)',
        'skills': r'(?i)(skills|technical\s+skills|core\s+competencies|expertise)',
        'experience': r'(?i)(experience|work\s+experience|professional\s+experience|employment)',
        'education': r'(?i)(education|academic|qualifications)',
        'projects': r'(?i)(projects|personal\s+projects)',
        'certifications': r'(?i)(certifications|certificates|accreditations)',
        'languages': r'(?i)(languages|language\s+proficiency)',
        'interests': r'(?i)(interests|hobbies)'
    }
    
    # Split text into lines for processing
    lines = text.split('\n')
    
    # Initialize sections dictionary
    sections = {
        'contact': [],
        'summary': [],
        'skills': [],
        'experience': [],
        'education': [],
        'projects': [],
        'certifications': [],
        'languages': [],
        'interests': [],
        'other': []  # For content that doesn't fit into known sections
    }
    
    # Initialize current section
    current_section = 'other'
    
    # Process each line
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check if this line is a section header
        found_section = False
        for section, pattern in section_patterns.items():
            if re.match(pattern, line, re.IGNORECASE) and len(line) < 50:  # Simple length check to avoid false positives
                current_section = section
                found_section = True
                break
        
        if not found_section:
            sections[current_section].append(line)
    
    return sections

def extract_contact_info(contact_section):
    """
    Extract structured contact information from the contact section.
    
    Args:
        contact_section (list): List of text lines from the contact section
        
    Returns:
        dict: Structured contact information
    """
    contact_info = {
        'name': '',
        'email': '',
        'phone': '',
        'location': '',
        'linkedin': '',
        'website': ''
    }
    
    # Join all lines
    text = ' '.join(contact_section)
    
    # Extract email
    email_match = re.search(r'[\w\.-]+@[\w\.-]+', text)
    if email_match:
        contact_info['email'] = email_match.group(0)
    
    # Extract phone (various formats)
    phone_match = re.search(r'(?:\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
    if phone_match:
        contact_info['phone'] = phone_match.group(0)
    
    # Extract LinkedIn URL
    linkedin_match = re.search(r'linkedin\.com/in/[\w-]+', text)
    if linkedin_match:
        contact_info['linkedin'] = 'https://' + linkedin_match.group(0)
    
    # Extract website
    website_match = re.search(r'https?://(?:www\.)?[\w\.-]+\.\w+', text)
    if website_match and 'linkedin.com' not in website_match.group(0):
        contact_info['website'] = website_match.group(0)
    
    # Assume the first line might contain the name if not too long
    if contact_section and len(contact_section[0]) < 50:
        contact_info['name'] = contact_section[0]
    
    # Try to extract location (this is approximate)
    location_patterns = [
        r'(?:located\s+in|location:?\s+)([^,\.]+(?:,\s*[^,\.]+)?)',
        r'([A-Z][a-zA-Z]+(?:[\s,]+[A-Z][a-zA-Z]+)+(?:[\s,]+[A-Z]{2})?)(?:\s*\d{5})?'
    ]
    
    for pattern in location_patterns:
        location_match = re.search(pattern, text)
        if location_match:
            contact_info['location'] = location_match.group(1).strip()
            break
    
    return contact_info

def extract_skills(skills_section):
    """
    Extract skills from the skills section.
    
    Args:
        skills_section (list): List of text lines from the skills section
        
    Returns:
        list: Extracted skills
    """
    skills = []
    
    # Join all lines
    text = ' '.join(skills_section)
    
    # Look for skill lists separated by common delimiters
    skill_lists = re.split(r'(?:\||•|,|;|\n)', text)
    
    for skill in skill_lists:
        skill = skill.strip()
        if skill and len(skill) < 50:  # Simple filter to avoid sentences
            skills.append(skill)
    
    return skills

def extract_experience(experience_section):
    """
    Extract work experience with improved structure detection.
    """
    experiences = []
    current_job = None
    current_role = None
    date_pattern = r'(?:\d{1,2}/\d{1,2}|\d{1,2}/\d{4}|\d{4}|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)'
    
    # If we have a large block of text, let's try to split it into sections first
    if len(experience_section) > 0 and any('\n' in line for line in experience_section):
        # Flatten and re-split by newlines
        full_text = '\n'.join(experience_section)
        experience_section = [line.strip() for line in full_text.split('\n') if line.strip()]
    
    i = 0
    while i < len(experience_section):
        line = experience_section[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
            
        # Check if this could be a job title (short capitalized text, no punctuation)
        if len(line) < 50 and line[0].isupper() and not any(x in line for x in ['.', ':', ',', ';']):
            # Save previous job if we have one
            if current_job:
                experiences.append(current_job)
                
            # Start a new job object
            current_job = {
                'title': line,
                'company': '',
                'location': '',
                'date_range': '',
                'description': []
            }
            i += 1
            
            # Check if next line might be company name
            if i < len(experience_section):
                company_line = experience_section[i].strip()
                if company_line and len(company_line) < 50:
                    current_job['company'] = company_line
                    i += 1
                    
                    # Check if next line might be location
                    if i < len(experience_section):
                        location_line = experience_section[i].strip()
                        # Location typically has a state/province code
                        if location_line and len(location_line) < 50 and re.search(r'[A-Z]{2}', location_line):
                            current_job['location'] = location_line
                            i += 1
                
                    # Check if next line might be date range
                    if i < len(experience_section):
                        date_line = experience_section[i].strip()
                        if date_line and re.search(date_pattern, date_line):
                            current_job['date_range'] = date_line
                            i += 1
        
        # If we have a job and this isn't a new job title, it's probably a bullet point
        elif current_job:
            # This is a description/bullet point for the current job
            current_job['description'].append(line)
            i += 1
        else:
            # If we can't categorize this line, just move on
            i += 1
    
    # Don't forget to add the last job
    if current_job:
        experiences.append(current_job)
    
    return experiences

def extract_education(education_section):
    """
    Extract education information from the education section.
    
    Args:
        education_section (list): List of text lines from the education section
        
    Returns:
        list: Structured education entries
    """
    education = []
    current_education = None
    
    for line in education_section:
        # Try to detect new education entry
        date_pattern = r'(?:\d{1,2}/\d{1,2}|\d{1,2}/\d{4}|\d{4}|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)'
        is_new_entry = (
            bool(re.search(rf'{date_pattern}.*{date_pattern}|{date_pattern}.*present|{date_pattern}.*current', line, re.IGNORECASE)) or
            bool(re.search(r'^[A-Z][^,\.]{0,50}$', line)) or  # Capitalized short line
            bool(re.search(r'(?:University|College|Institute|School)', line))
        )
        
        if is_new_entry:
            # Save previous entry if exists
            if current_education:
                education.append(current_education)
            
            # Initialize new education entry
            current_education = {
                'institution': '',
                'degree': '',
                'date_range': '',
                'details': []
            }
            
            # Try to extract date range
            date_match = re.search(rf'({date_pattern}.*?{date_pattern}|{date_pattern}.*?present|{date_pattern}.*?current)', line, re.IGNORECASE)
            if date_match:
                current_education['date_range'] = date_match.group(0)
                # Remove date from line for further processing
                line = re.sub(rf'{re.escape(date_match.group(0))}', '', line).strip()
            
            # Check for degree information
            degree_patterns = [
                r'(?:Bachelor|Master|PhD|Doctorate|B\.S\.|M\.S\.|B\.A\.|M\.B\.A\.|Ph\.D\.)[^,\.]*',
                r'(?:BS|MS|BA|MBA|PhD)[^,\.]*'
            ]
            
            for pattern in degree_patterns:
                degree_match = re.search(pattern, line, re.IGNORECASE)
                if degree_match:
                    current_education['degree'] = degree_match.group(0).strip()
                    # Remove degree from line
                    line = re.sub(rf'{re.escape(degree_match.group(0))}', '', line).strip()
                    break
            
            # Assume what's left is the institution
            if line:
                current_education['institution'] = line
                
        elif current_education:
            # Add line to details of current education
            if line.strip():
                current_education['details'].append(line.strip())
    
    # Add the last education entry
    if current_education:
        education.append(current_education)
    
    return education

def parse_resume(docx_path):
    """
    Main function to parse a resume DOCX and extract structured data.
    
    Args:
        docx_path (str): Path to the resume DOCX file
        
    Returns:
        dict: Structured resume data in JSON format
    """
    try:
        # Extract text from DOCX
        resume_text = extract_text_from_docx(docx_path)
        
        # Identify sections
        sections = identify_sections(resume_text)
        
        # Extract structured information from each section
        contact_info = extract_contact_info(sections['contact'])
        skills = extract_skills(sections['skills'])
        experience = extract_experience(sections['experience'])
        education = extract_education(sections['education'])
        
        # Combine into a structured resume object
        resume_data = {
            'contact_info': contact_info,
            'summary': '\n'.join(sections['summary']),
            'skills': skills,
            'experience': experience,
            'education': education,
            'projects': '\n'.join(sections['projects']),
            'certifications': '\n'.join(sections['certifications']),
            'languages': '\n'.join(sections['languages']),
            'interests': '\n'.join(sections['interests']),
            'other': '\n'.join(sections['other']),
            'raw_text': resume_text  # Include raw text for AI processing
        }
        
        return resume_data
        
    except Exception as e:
        raise Exception(f"Error parsing resume: {str(e)}")
