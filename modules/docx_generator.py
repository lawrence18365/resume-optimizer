"""
Professional Resume DOCX Generator - Creates beautifully formatted resume documents
"""
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_bottom_border(paragraph):
    """Add a bottom border to a paragraph."""
    p = paragraph._p  # p is the paragraph element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.append(pBdr)
    
    # Add bottom border
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')  # Border size
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4472C4')  # Border color - professional blue
    pBdr.append(bottom)

def generate_docx(optimized_resume, output_path):
    """Generate a beautifully formatted professional DOCX resume."""
    doc = docx.Document()
    
    # ===== DOCUMENT SETUP =====
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # ===== STYLE DEFINITIONS =====
    styles = doc.styles
    
    # Name style - Large, bold, centered
    name_style = styles.add_style('Name', WD_STYLE_TYPE.PARAGRAPH)
    name_font = name_style.font
    name_font.name = 'Calibri'
    name_font.size = Pt(20)
    name_font.bold = True
    name_font.color.rgb = RGBColor(0, 51, 102)  # Dark blue for professional look
    name_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_style.paragraph_format.space_after = Pt(2)
    
    # Contact style - Centered contact information
    contact_style = styles.add_style('Contact', WD_STYLE_TYPE.PARAGRAPH)
    contact_font = contact_style.font
    contact_font.name = 'Calibri'
    contact_font.size = Pt(10)
    contact_font.color.rgb = RGBColor(0, 0, 0)
    contact_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_style.paragraph_format.space_after = Pt(12)
    
    # Section Header style - Bold with bottom border
    section_style = styles.add_style('Section', WD_STYLE_TYPE.PARAGRAPH)
    section_font = section_style.font
    section_font.name = 'Calibri'
    section_font.size = Pt(14)
    section_font.bold = True
    section_font.color.rgb = RGBColor(0, 51, 102)  # Dark blue for professional look
    section_style.paragraph_format.space_before = Pt(15)
    section_style.paragraph_format.space_after = Pt(6)
    section_style.paragraph_format.keep_with_next = True
    
    # Job Title style - Bold
    job_style = styles.add_style('JobTitle', WD_STYLE_TYPE.PARAGRAPH)
    job_font = job_style.font
    job_font.name = 'Calibri'
    job_font.size = Pt(12)
    job_font.bold = True
    job_style.paragraph_format.space_before = Pt(14)
    job_style.paragraph_format.space_after = Pt(0)
    job_style.paragraph_format.keep_with_next = True
    
    # Company style - Italic with location
    company_style = styles.add_style('Company', WD_STYLE_TYPE.PARAGRAPH)
    company_font = company_style.font
    company_font.name = 'Calibri'
    company_font.size = Pt(11)
    company_font.italic = True
    company_style.paragraph_format.space_before = Pt(2)
    company_style.paragraph_format.space_after = Pt(0)
    company_style.paragraph_format.keep_with_next = True
    
    # Date style - Right-aligned
    date_style = styles.add_style('Date', WD_STYLE_TYPE.PARAGRAPH)
    date_font = date_style.font
    date_font.name = 'Calibri'
    date_font.size = Pt(10)
    date_font.italic = True
    date_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    date_style.paragraph_format.space_before = Pt(0)
    date_style.paragraph_format.space_after = Pt(4)
    
    # Bullet style - For experience bullet points
    bullet_style = styles.add_style('Bullet', WD_STYLE_TYPE.PARAGRAPH)
    bullet_font = bullet_style.font
    bullet_font.name = 'Calibri'
    bullet_font.size = Pt(10.5)
    bullet_style.paragraph_format.left_indent = Inches(0.25)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.15)
    bullet_style.paragraph_format.space_after = Pt(3)
    bullet_style.paragraph_format.space_before = Pt(0)
    bullet_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # Skill style - For skills section
    skill_style = styles.add_style('Skill', WD_STYLE_TYPE.PARAGRAPH)
    skill_font = skill_style.font
    skill_font.name = 'Calibri'
    skill_font.size = Pt(11)
    skill_style.paragraph_format.space_after = Pt(6)
    
    # Summary style - For professional summary
    summary_style = styles.add_style('Summary', WD_STYLE_TYPE.PARAGRAPH)
    summary_font = summary_style.font
    summary_font.name = 'Calibri'
    summary_font.size = Pt(11)
    summary_style.paragraph_format.space_after = Pt(6)
    
    # ===== CONTENT GENERATION =====
    
    # ----- CONTACT INFORMATION -----
    contact = optimized_resume.get('contact_info', {})
    name = contact.get('name', '')
    
    if name:
        doc.add_paragraph(name.upper(), style='Name')
    
    # Contact details
    contact_parts = []
    if contact.get('email'):
        contact_parts.append(contact['email'])
    if contact.get('phone'):
        contact_parts.append(contact['phone'])
    if contact.get('location'):
        contact_parts.append(contact['location'])
        
    if contact_parts:
        contact_para = doc.add_paragraph(style='Contact')
        contact_para.add_run(' | '.join(contact_parts))
    
    # Links (LinkedIn, website)
    link_parts = []
    if contact.get('linkedin'):
        link_parts.append(contact['linkedin'])
    if contact.get('website'):
        link_parts.append(contact['website'])
        
    if link_parts:
        links_para = doc.add_paragraph(style='Contact')
        links_para.add_run(' | '.join(link_parts))
    
    # ----- PROFESSIONAL SUMMARY -----
    summary = optimized_resume.get('summary', '')
    if summary:
        summary_header = doc.add_paragraph("PROFESSIONAL SUMMARY", style='Section')
        add_bottom_border(summary_header)
        
        # Split summary into paragraphs for better readability
        paragraphs = summary.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip(), style='Summary')
    
    # ----- SKILLS -----
    skills = optimized_resume.get('skills', [])
    if skills:
        skills_header = doc.add_paragraph("SKILLS", style='Section')
        add_bottom_border(skills_header)
        
        # Format skills as a readable list
        if isinstance(skills, list):
            skill_text = ', '.join(skills)
        else:
            skill_text = skills
            
        doc.add_paragraph(skill_text, style='Skill')
    
    # ----- PROFESSIONAL EXPERIENCE -----
    experience = optimized_resume.get('experience', [])
    if experience:
        exp_header = doc.add_paragraph("PROFESSIONAL EXPERIENCE", style='Section')
        add_bottom_border(exp_header)
        
        for job in experience:
            # Job Title
            title = job.get('title', '')
            if title:
                job_para = doc.add_paragraph(title, style='JobTitle')
            
            # Company and Location
            company = job.get('company', '')
            location = job.get('location', '')
            company_text = company
            if location:
                company_text += f" | {location}"
                
            if company_text:
                doc.add_paragraph(company_text, style='Company')
            
            # Date Range - Right aligned
            date_range = job.get('date_range', '')
            if date_range:
                doc.add_paragraph(date_range, style='Date')
            
            # Description bullets
            descriptions = job.get('description', [])
            if descriptions:
                if isinstance(descriptions, list):
                    for desc in descriptions:
                        if desc.strip():
                            bullet_para = doc.add_paragraph(style='Bullet')
                            bullet_para.add_run("• ").bold = True
                            bullet_para.add_run(desc.strip())
                else:
                    for line in descriptions.split('\n'):
                        if line.strip():
                            bullet_para = doc.add_paragraph(style='Bullet')
                            bullet_para.add_run("• ").bold = True
                            bullet_para.add_run(line.strip())
    
    # ----- EDUCATION -----
    education = optimized_resume.get('education', [])
    if education:
        edu_header = doc.add_paragraph("EDUCATION", style='Section')
        add_bottom_border(edu_header)
        
        for edu in education:
            degree = edu.get('degree', '')
            institution = edu.get('institution', '')
            
            edu_line = []
            if degree:
                edu_line.append(degree)
            if institution:
                edu_line.append(institution)
                
            if edu_line:
                doc.add_paragraph(' - '.join(edu_line), style='JobTitle')
            
            # Date Range
            date_range = edu.get('date_range', '')
            if date_range:
                doc.add_paragraph(date_range, style='Date')
            
            # Details
            details = edu.get('details', [])
            if details:
                if isinstance(details, list):
                    for detail in details:
                        if detail.strip():
                            bullet_para = doc.add_paragraph(style='Bullet')
                            bullet_para.add_run("• ").bold = True
                            bullet_para.add_run(detail.strip())
                else:
                    for line in details.split('\n'):
                        if line.strip():
                            bullet_para = doc.add_paragraph(style='Bullet')
                            bullet_para.add_run("• ").bold = True
                            bullet_para.add_run(line.strip())
                            
    # ----- CERTIFICATIONS (if present) -----
    certifications = optimized_resume.get('certifications', '')
    if certifications:
        cert_header = doc.add_paragraph("CERTIFICATIONS", style='Section')
        add_bottom_border(cert_header)
        
        if isinstance(certifications, list):
            for cert in certifications:
                if cert.strip():
                    bullet_para = doc.add_paragraph(style='Bullet')
                    bullet_para.add_run("• ").bold = True
                    bullet_para.add_run(cert.strip())
        else:
            for line in certifications.split('\n'):
                if line.strip():
                    bullet_para = doc.add_paragraph(style='Bullet')
                    bullet_para.add_run("• ").bold = True
                    bullet_para.add_run(line.strip())
    
    # ----- PROJECTS (if present) -----
    projects = optimized_resume.get('projects', '')
    if projects:
        proj_header = doc.add_paragraph("PROJECTS", style='Section')
        add_bottom_border(proj_header)
        
        if isinstance(projects, list):
            for proj in projects:
                if proj.strip():
                    bullet_para = doc.add_paragraph(style='Bullet')
                    bullet_para.add_run("• ").bold = True
                    bullet_para.add_run(proj.strip())
        else:
            for line in projects.split('\n'):
                if line.strip():
                    bullet_para = doc.add_paragraph(style='Bullet')
                    bullet_para.add_run("• ").bold = True
                    bullet_para.add_run(line.strip())
                    
    # Save the document
    doc.save(output_path)
    return doc